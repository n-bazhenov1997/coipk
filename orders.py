# ============================================================
# Импорты
# ============================================================

# Стандартная библиотека
from datetime import date
from pathlib import Path

# Сторонние библиотеки
import xlwings as xw
from docx import Document

# GUI
from tkinter import messagebox

# Локальные модули
from scripts import check_selection


# ============================================================
# 1️⃣ Проверка условий для запуска программы
# ============================================================
def check_conditions() -> bool:
    """
    Проверяет, можно ли запускать программу на основе выделенного диапазона Excel.

    Проверяет:
        - Существует ли выделение в активной книге.
        - Все ли номера приказов в выделенном диапазоне одинаковы.

    Возвращает:
        bool: True, если условия выполнены, False иначе.
    """
    try:
        book = xw.books.active
        sel = book.selection

        # Проверка корректности выделения с помощью локального скрипта
        if check_selection:
            return False

        # Проверяем, что все номера приказов одинаковы
        for i in range(sel.rows.count):
            if i == 0:
                ord_num = sel[0, 2].value  # Берём номер приказа из первой строки
            elif sel[i, 2].value != ord_num:
                messagebox.showerror(
                    'Ошибка',
                    'Один из номеров приказа в выделенном диапазоне не соответствует другим номерам'
                )
                return False
        else:
            return True
    except Exception as e:
        messagebox.showerror('Ошибка', f'Ошибка при проверке условий: {e}')
        return False


# ============================================================
# 2️⃣ Сбор информации из выделенного диапазона Excel
# ============================================================
def collect_info() -> tuple[dict[str, str | date | int], list[str]]:
    """
    Собирает информацию для приказа из выделенного диапазона Excel.
    
    Предполагается, что проверка условий уже выполнена через `check_conditions`.

    Возвращает:
        main_info (dict): Словарь с основными данными приказа.
        studs_names (list): Список полных имен студентов.
    """
    book = xw.books.active
    sel = book.selection

    main_info = {
        'ord_date': sel[0, 0].value,    # Дата приказа
        'ord_num': sel[0, 2].value,     # Номер приказа
        'prog': sel[0, 4].value,        # Название программы
        'ord_type': sel[0, 9].value,    # Тип приказа
        'prot_num': sel[0, 10].value,   # Номер протокола
        'teacher': sel[0, 12].value,    # ФИО преподавателя
        'start_date': sel[0, 13].value, # Дата начала обучения
        'end_date': sel[0, 14].value,   # Дата окончания обучения
        'hours': sel[0, 15].value       # Количество часов
    }

    studs_names = [sel[i, 11].value for i in range(sel.rows.count)]

    return main_info, studs_names


# ============================================================
# 3️⃣ Определение шаблона для приказа
# ============================================================
def def_template(ord_num: str, ord_type: str) -> Path | None:
    """
    Определяет путь к шаблону приказа на основе номера и типа приказа.

    Аргументы:
        ord_num (str): Номер приказа.
        ord_type (str): Тип приказа ('зачисление' или 'отчисление').

    Возвращает:
        Path | None: Полный путь к файлу шаблона или None, если тип или шаблон не найден.
    """
    templates_folder = Path(r'\\Win-1pvk6u4un23\общая папка\Macros\Templates\orders')

    doctype = {
        'ДУ': 'upgrade',
        'ДД': 'diploma',
        'ДС': 'vocational'
    }

    template_type = {
        'зачисление': 'enroll.docx',
        'отчисление': 'complete.docx'
    }

    subfolder = doctype.get(ord_num[:2])
    filename = template_type.get(ord_type)

    if subfolder is None:
        messagebox.showerror('Ошибка', 'Не найден тип документа по переданному номеру приказа')
        return None
    elif filename is None:
        messagebox.showerror('Ошибка', 'Не найден шаблон приказа по переданному виду приказа')
        return None
    else:
        return templates_folder / subfolder / filename


# ============================================================
# 4️⃣ Внесение информации в шаблон
# ============================================================
def fill_template(template_path: Path, main_info: dict[str | date | int], studs_names: list[str]):
    """
    Заполняет шаблон приказа данными из Excel.

    Аргументы:
        template_path (Path): Путь к файлу шаблона.
        main_info (dict): Основная информация о приказе.
        studs_names (list): Список студентов.
    """
    try:
        doc = Document(template_path)

        # Плейсхолдеры для замены текста
        data = {
            '{{ord_date}}': main_info['ord_date'],
            '{{ord_num}}': main_info['ord_num'],
            '{{prog}}': main_info['prog'],
            '{{start_date}}': main_info['start_date'],
            '{{end_date}}': main_info['end_date'],
            '{{hours}}': main_info['hours'],
        }

        # Замена плейсхолдеров на значения
        for p in doc.paragraphs:
            for placeholder, value in data.items():
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

        # Заполнение основной таблицы студентами
        main_table = doc.tables[0]
        for i, name in enumerate(studs_names, start=1):
            row = main_table.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = str(name)

        # Сохраняем заполненный документ под новым именем
        output_path = template_path.with_name(f'filled_{template_path.name}')
        doc.save(output_path)

        messagebox.showinfo('Готово', f'Приказ успешно сформирован: {output_path.name}')
    except Exception as e:
        messagebox.showerror('Ошибка', f'Ошибка при заполнении шаблона: {e}')


# ============================================================
# 5️⃣ Основная функция
# ============================================================
def main():
    """
    Основной поток программы:
        1. Проверяет условия для запуска.
        2. Собирает информацию из Excel.
        3. Определяет путь к шаблону.
        4. Заполняет шаблон данными.
    """
    if not check_conditions():
        return

    main_info, studs_names = collect_info()

    template_path = def_template(main_info['ord_num'], main_info['ord_type'])
    if template_path is None:
        return

    fill_template(template_path, main_info, studs_names)